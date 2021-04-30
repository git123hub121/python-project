#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from faker import Faker
import random
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import font_manager
from docx import Document
from docx.oxml.ns import qn
import numpy as np

myfont = font_manager.FontProperties(fname=r"C:\Windows\Fonts\simsun.ttc")
# 伪造数据
fake = Faker('zh_CN')  
name = []
sex= []
score1 = []
score2 = []
score3 = []
score4 = []
number = range(1,31)
for _ in range(30):
    name.append(fake.simple_profile(sex=None)['name'])
    sex.append(fake.simple_profile(sex=None)['sex'])
    score1.append(random.randint(40,100))
    score2.append(random.randint(40,100))
    score3.append(random.randint(40,100))
    score4.append(random.randint(200,300))
    
# 写入Excel
df = pd.DataFrame({
        '学号':number,
        '姓名':name,    
        '性别':sex,
        '语文':score1,
        '数学':score2,
        '英语':score3,
        '理综':score4
        })

df = df.set_index('学号')

df.to_excel('Part3_学生成绩单.xlsx')


## 读取数据
students = pd.read_excel('Part3_学生成绩单.xlsx')

# 排序名单
students['总分'] = students.语文 + students.数学 + students.英语 + students.理综
students.sort_values(by='总分', inplace=True, ascending=False)
students.reset_index(drop=True, inplace=True)

#学生成绩汇总表
ax = students.plot.bar(x='姓名', y=['语文','数学','英语','理综'], stacked=True)
plt.title('学生成绩汇总图', fontsize=16,fontproperties=myfont)   # fontproperties=font
plt.xlabel('姓名',  fontsize=10,fontproperties=myfont)  # fontproperties=font,
plt.xticks(rotation='45', fontsize=8,fontproperties=myfont)  #  fontproperties=font, 
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
plt.tight_layout()
plt.savefig('Part3_data.jpg')
plt.show()


students_score =students[['语文','数学','英语','理综']]
scores = {col:students_score[col].tolist() for col in students_score.columns} 

from itertools import groupby
# 绘制图表
plt.rcParams['font.sans-serif']= [r'SimHei'] # 解决图例中文乱码
plt.rcParams['axes.unicode_minus']=False


def splitScore(score):
    if score >=85:
        return '优'
    elif score >=60:
        return '及格'
    else :
        return '不及格'
    
def splitScore_lizong(score):
    if score >=260:
        return '优'
    elif score >=240:
        return '及格'
    else :
        return '不及格'
    
def pie_plot(scores):
    ratios=dict()
    for subject,subjectScore in scores.items():
        ratios[subject]={}
        if subject !='理综':
            for category,num in groupby(sorted(subjectScore),splitScore):
                ratios[subject][category]= len(tuple(num))
        else:
            for category,num in groupby(sorted(subjectScore),splitScore_lizong):
                ratios[subject][category]= len(tuple(num))
            
    
    
    fig ,axs = plt.subplots(2,2) 
    # 画子图
    axs.shape=1,4
    for index,subjectData in enumerate(ratios.items()):
        plt.sca(axs[0][index])
        subjectName,subjectRatio = subjectData
        plt.pie(list(subjectRatio.values()),labels=list(subjectRatio.keys()),autopct='%1.1f%%')
        plt.xlabel(subjectName)
        plt.legend(loc="right",bbox_to_anchor=(1, 0, 0.5, 1))
    
    plt.savefig('Part4_data.jpg')
    plt.show()
    
    
pie_plot(scores)


# 操作Word
document = Document()
document.add_heading('本次考试学生成绩分析报告', level=0)
first_student = students.iloc[0,:]['姓名']
first_score = students.iloc[0,:]['总分']
# 设置格式
document.styles['Normal'].font.name = 'Times New Roman'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

p = document.add_paragraph('本次测评，全班共有{}名同学参加考试，其中分数总分排名第一的同学是'.format(len(students.姓名)),style='Heading 3')
p.add_run(str(first_student)).bold = True
p.add_run('，分数为')
p.add_run(str(first_score)).bold = True
p.add_run('.学生考试总体成绩如下')

table = document.add_table(rows=len(students.姓名)+1, cols=6, style='Medium Shading 1 Accent 5')
table.cell(0,0).text = '姓名'
table.cell(0,1).text = '语文'
table.cell(0,2).text = '数学'
table.cell(0,3).text = '英语'
table.cell(0,4).text = '理综'
table.cell(0,5).text = '总分'

for i,(index,row) in enumerate(students.iterrows()):
    table.cell(i+1, 0).text = str(row['姓名'])
    table.cell(i+1, 1).text = str(row['语文'])
    table.cell(i+1, 2).text = str(row['数学'])
    table.cell(i+1, 3).text = str(row['英语'])
    table.cell(i+1, 4).text = str(row['理综'])
    table.cell(i+1, 5).text = str(row['总分'])
    

p = document.add_paragraph('成绩排行榜',style='Heading 3')
#p.add_run('成绩排行榜')  
document.add_picture('Part3_data.jpg')

p = document.add_paragraph('各科成绩分布',style='Heading 3')
#p.add_run('成绩排行榜')  
document.add_picture('Part4_data.jpg')


document.add_paragraph('王欢的成绩概括分析',style='Heading 3')


# 伪造数据
fake = Faker('zh_CN')  
name = []
sex= []
score1 = []
score2 = []
score3 = []
score4 = []
number = range(1,31)
for _ in range(30):
    name.append(fake.simple_profile(sex=None)['name'])
    sex.append(fake.simple_profile(sex=None)['sex'])
    score1.append(random.randint(80,100))
    score2.append(random.randint(90,100))
    score3.append(random.randint(70,100))
    score4.append(random.randint(230,300))
    
# 写入Excel
temp_df = pd.DataFrame({
        '历次考试':number,
        '姓名':'王欢',    
        '语文':score1,
        '数学':score2,
        '英语':score3,
        '理综':score4
        })


df = temp_df.describe()
df = df.reset_index()
df.rename(columns={'index':'衡量指标'} ,inplace=True)
print(df)

#document = Document()

table = document.add_table(rows=len(df.历次考试)+1, cols=5, style='Medium Shading 1 Accent 5')
table.cell(0,0).text = '衡量指标'
table.cell(0,1).text = '语文'
table.cell(0,2).text = '数学'
table.cell(0,3).text = '英语'
table.cell(0,4).text = '理综'


for i,(index,row) in enumerate(df.iterrows()):
    table.cell(i+1, 0).text = str(row['衡量指标'])
    table.cell(i+1, 1).text = str(row['语文'])
    table.cell(i+1, 2).text = str(row['数学'])
    table.cell(i+1, 3).text = str(row['英语'])
    table.cell(i+1, 4).text = str(row['理综'])



# Fixing random state for reproducibility
np.random.seed(19680801)

plt.subplot(411)
plt.plot(temp_df['语文'],'y*-',label="语文")

plt.subplot(412)
plt.plot(temp_df['数学'],'b*-')
plt.subplot(413)
plt.plot(temp_df['英语'],'r*-')
plt.subplot(414)
plt.plot(temp_df['理综'])
plt.savefig('Part6_data.jpg')
plt.show()


document.add_paragraph('王欢的历次考试趋势图',style='Heading 3')
document.add_picture('Part6_data.jpg')

document.save('Part3_学生成绩分析报告.docx')

#项目很好，但是有点看不懂，有时间找一下源数据
















